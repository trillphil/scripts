#!/usr/bin/env python3

import docx
import re
import argparse
from colorama import Fore, Style, init as colorinit


__version__ = "0.1"

# Add common template strings to identify here
common_strings = [
'CLIENTLONG',
'CLIENTSHORT',
'TODO',
'CHANGE',
'(Name)',
'244-xxxx',
'xxx@swordshield.com',
'a X.509',
'https://google.com',
]

# Add strings within parenthesis that are NOT acronyms that need to be in Appendix A
#  many of these are template strings, certifications, or in the methodology sections
acronym_false_positives = [
'(compromise)',
'(Gauntlet)',
'(hackers)',
'(trusted)',
'(downloads)',
'(s)',
'(expired)',
'(Medium)',
'(OSCP)',
'(GIAC)',
'(GSEC)',
'(GCIA)',
'(GCIH)',
'(GWAPT)',
'(CISSP)',
'(Redacted)',
'(Name)',
'(CLIENTSHORT)',
]


def print_header():
    """Prints tool art"""
    
    art = """%s
######################################
#       SA Report Checker v%s       #
######################################
""" % (Fore.CYAN,__version__)
    
    print(art)
    
    
def find_acronyms(report, verbose):
    """Attempts to find acronyms, repeated acronyms, and output a list for documentation
    
    Arguments:
    report - the docx report object that will be searched    
    verbose - may print extra information if set
    
    Assumptions:
    Acronyms: this function assumes that anything within parenthesis is an acronym
    """
    
    acronyms = []
    acronyms_dict = {}
    regex = r"(\([a-zA-Z]+\))"

    print("\n[*] Checking for acronyms")
    
    for para in report.paragraphs:
        matches = re.findall(regex, para.text)
        if matches:
            for match in matches:
                # skip false positives
                if match in acronym_false_positives:
                    continue
                elif match in acronyms:
                    print("%s[!] \"%s\" already defined!" % (Fore.YELLOW, match))
                    
                    if verbose:
                        print("[*] Check near:\n  %s" % (para.text))
                else:
                    # turn the paragraph into an array so we can reference the acronym's
                    #  relative position within it
                    words = para.text.replace('.', '').replace(',', '').split(' ')
                    pos = words.index(match)
                    key = match.replace('(', '').replace(')', '')
                    
                    # if the word prior to the acronym is hyphenated, then the acryonym
                    #  likely points to it. Otherwise, we'll go off of white space
                    if words[pos - 1].count('-') > 0:
                        value = words[pos - 1]
                    else:
                        value = " ".join(words[pos - len(key):pos])
                    acronyms.append(match)
                    acronyms_dict[key] = value

    print('%s\n[*] The following acronyms may need to be added to the Appendix:' % (Fore.GREEN))

    # sort acronym keys alphabetically
    acronyms = sorted(acronyms_dict.keys())

    # print acronym keys and values
    for key in acronyms:
        print("%s\t\t%s" % (key, acronyms_dict[key]))
        
        
def find_common_oversights(report, verbose):
    """ Finds commonly overlooked template strings
    
    Arguments:
    report - the docx report object that will be searched
    verbose - may print extra information if set
    """
    
    query = ""
    regex = ""
    found_strings = {}
    
    print("\n[*] Checking for %s" % (", ".join(common_strings)))
    
    
    # build regex query and dictionary to store strings with their counts
    for i in common_strings:
        query += i + '|'
        found_strings[i] = 0
    
    # suing replace instead of re.escape because of the | chars
    query = query.replace('(', '\(').replace(')','\)').replace('.', '\.').replace('@','\@')
    regex = r"({})".format(query[0:-1])
    
    # search every paragraph object for the common strings
    for para in report.paragraphs:
        matches = re.findall(regex, para.text)
        
        # if a common string is found, increment the dictionary value for that strings
        if matches:
            for match in matches:
                found_strings[match] += 1
    
    # print strings that were found in the documents and the count for each
    for key, value in found_strings.items():
        if value:
            print("%s[!] Found \"%s\" %s times!" % (Fore.YELLOW, key, value))


def find_references(report, verbose):
    """Checks table and figure references against number of tables and figures
    
    Arguments:
    report - the docx report object that will be searched
    verbose - may print extra information if set
    
    Assumptions:
    Captions - must be in form Table XX: and Figure XX:. if the caption does not
      contain a colon this function will not work properly
    """
    
    tables = []
    figures = []
    references = []
    regex = r"(Table [\d]{1,3}.|Figure [\d]{1,3}.)"
    
    print("\n[*] Checking for table and figure references")
        
    for para in report.paragraphs:
        matches = re.findall(regex, para.text)
        
        if matches:
            for match in matches:
                # if the match ends with a colon, it's a caption
                if match.endswith(":"):
                    if "Table" in match and match not in tables:
                        tables.append(match[0:-1])
                    elif "Figure" in match and match not in figures:
                        figures.append(match[0:-1])
                # otherwise, it's a reference
                else:
                    if not match[-1].isdigit():
                        references.append(match[0:-1].strip())
                    else:
                        references.append(match.strip())
    
    
    for table in tables:
        if table not in references:
            print("%s[!] Reference missing for %s" % (Fore.YELLOW, table))
    
    for figure in figures:
        if figure not in references:
            print("%s[!] Reference missing for %s" % (Fore.YELLOW, figure))


def main():
    """the main function"""
    
    parser = argparse.ArgumentParser(description='SA Report Checker', 
                    formatter_class=argparse.RawTextHelpFormatter)
    parser.add_argument('--version', action='version',
                    version='%(prog)s {version}'.format(version=__version__))
    parser.add_argument("report", help="the report to check")
    parser.add_argument("-v", "--verbose", help="enable verbose output", action="store_true")
    parser.add_argument("-c", "--checks" , help="checks to run against report\n"
            "ALL = all checks (default)\n"
            "o   = common oversights\n"
            "r   = caption references\n"
            "a   = acronyms\n"
            ,  default="ALL")
    args = parser.parse_args()
    
    # set variables from CLI args
    report = docx.Document(args.report)
    verbose = args.verbose
    checks = args.checks
    
    # set colorama to automatically reset after each print
    colorinit(autoreset=True)
    
    
    print_header()
    
    if checks == "ALL":
        find_common_oversights(report, verbose)
        find_references(report, verbose)
        find_acronyms(report, verbose)
    else:
        if "o" in checks:
            find_common_oversights(report, verbose)
        if "r" in checks:
            find_references(report, verbose)
        if "a" in checks:
            find_acronyms(report, verbose)

if __name__ == '__main__':
    main()
